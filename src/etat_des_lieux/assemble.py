"""Markdown etat des lieux -> DOCX.

Placeholder converter: turns # / ## / ### headings into Word headings and
everything else into paragraphs. Will be replaced with template-based assembly
once the user provides a DOCX template.
"""
from pathlib import Path
from docx import Document


def to_docx(markdown: str, out_path: Path) -> Path:
    doc = Document()
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
